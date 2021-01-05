import mimetypes
import os
import uuid
from datetime import datetime
from tempfile import NamedTemporaryFile

import docx
from celery.utils.log import get_task_logger
from flask import current_app

from app.celery import ContextTask
from app.extensions import celery
from app.models import Document as DocumentModel, User as UserModel
from app.utils import to_readable
from app.utils.constants import PDF_MIME_TYPE, MS_WORD_MIME_TYPE
from app.utils.file_storage import FileStorage
from app.utils.libreoffice import convert_to
from app.utils.request_query_operator import RequestQueryOperator as rqo
from app.serializers import DocumentSerializer, UserSerializer

logger = get_task_logger(__name__)

_COLUMN_DISPLAY_ORDER = ['name', 'last_name', 'email', 'birth_date', 'role',
                         'created_at', 'updated_at', 'deleted_at']


def _add_table_user_data(users_query: list, rows: list) -> None:
    user_list = []

    for user in users_query:
        user_dict = {
            'role': user.get('roles')[0].get('name'),
        }
        del user['roles']

        user_dict.update({
            k: to_readable(v)
            for (k, v) in user.items()
            if k in _COLUMN_DISPLAY_ORDER
        })

        user_dict = dict(
            sorted(user_dict.items(),
                   key=lambda x: _COLUMN_DISPLAY_ORDER.index(x[0]))
        )
        user_list.append(user_dict)

    for i, user_dict in enumerate(user_list):
        user_values = list(user_dict.values())
        rows.append(user_values)


def _add_table_column_names(rows: list, original_column_names: set) -> None:
    formatted_column_names = [
        column.title().replace('_', ' ')
        for column in original_column_names
        if column
    ]

    rows.append(formatted_column_names)


def _get_user_data(request_data: dict) -> list:
    page_number, items_per_page, order_by = rqo.get_request_query_fields(UserModel,
                                                                         request_data)

    query = UserModel.select()
    query = rqo.create_search_query(UserModel, query, request_data)
    query = (query.order_by(*order_by)
             .paginate(page_number, items_per_page))

    user_serializer = UserSerializer(many=True)
    user_list = user_serializer.dump(list(query))

    return user_list


@celery.task(bind=True, base=ContextTask)
def export_user_data_in_word_task(self, created_by: int, request_data: dict, to_pdf: int):
    def _write_docx_content(rows: list, document: docx.Document) -> None:
        header_fields = rows[0]
        assert len(header_fields) == len(_COLUMN_DISPLAY_ORDER)

        table = document.add_table(rows=len(rows), cols=len(header_fields))

        for i in range(len(rows)):
            row = table.rows[i]
            for j, table_cell in enumerate(rows[i]):
                row.cells[j].text = str(table_cell)

            self.update_state(state='PROGRESS', meta={
                'current': i,
                'total': self.total_progress,
                'status': 'In progress...',
            })

    user_list = _get_user_data(request_data)

    # Word table rows + 2 (Word table header and save data in database)
    self.total_progress = len(user_list) + 2
    tempfile_suffix = '.docx'
    tempfile = NamedTemporaryFile(suffix=tempfile_suffix)
    table_data = []
    mime_type = PDF_MIME_TYPE if to_pdf else MS_WORD_MIME_TYPE

    self.update_state(state='PROGRESS', meta={
        'current': 0,
        'total': self.total_progress,
        'status': 'In progress...',
    })

    _add_table_column_names(table_data, set(_COLUMN_DISPLAY_ORDER))
    _add_table_user_data(user_list, table_data)

    document = docx.Document()
    _write_docx_content(table_data, document)
    document.save(tempfile.name)

    directory_path = current_app.config.get('STORAGE_DIRECTORY')
    temp_filename = FileStorage.get_basename(tempfile.name, include_path=False)

    if to_pdf:
        convert_to(directory_path, tempfile.name)
    else:
        dst = f'{directory_path}/{temp_filename}{tempfile_suffix}'
        FileStorage.copy_file(tempfile.name, dst)

    file_extension = mimetypes.guess_extension(mime_type)
    internal_filename = '%s%s' % (uuid.uuid1().hex, file_extension)
    filepath = '%s/%s' % (directory_path, internal_filename)

    try:
        file_prefix = datetime.utcnow().strftime('%Y%m%d')
        basename = f'{file_prefix}_users'
        filename = f'{basename}{file_extension}'

        src = '%s/%s%s' % (directory_path, temp_filename, file_extension)
        FileStorage.rename(src, filepath)

        data = {
            'created_by': created_by,
            'name': filename,
            'internal_filename': internal_filename,
            'mime_type': mime_type,
            'directory_path': directory_path,
            'size': FileStorage.get_filesize(filepath),
        }

        document = DocumentModel.create(**data)
    except Exception as e:
        if os.path.exists(filepath):
            os.remove(filepath)
        logger.debug(e)
        raise e

    document_serializer = DocumentSerializer()
    document_data = document_serializer.dump(document)

    return {
        'current': self.total_progress,
        'total': self.total_progress,
        'status': 'Task completed!',
        'result': document_data,
    }
